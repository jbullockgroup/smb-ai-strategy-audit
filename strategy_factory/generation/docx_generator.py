"""
Word document generator for strategy report and statement of work.

Creates professional Word documents from synthesized deliverables.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..config import (
    OUTPUT_DIR,
    DELIVERABLES,
)

# Auditor branding logo for cover page footer
_AUDITOR_LOGO = Path(__file__).resolve().parent.parent.parent / "assets" / "auditor-logo" / "auditor-logo.png"
from ..models import (
    SynthesisOutput,
    ResearchOutput,
    CompanyInput,
)


class DocxGenerator:
    """
    Generates Word documents from deliverables.

    Responsibilities:
    - Create Final AI Strategy Report
    - Create Statement of Work / Engagement Letter
    - Apply consistent styling and formatting
    - Include relevant content from all deliverables
    """

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize the Word document generator.

        Args:
            output_dir: Base output directory.
        """
        self.output_dir = output_dir or OUTPUT_DIR

    def generate_strategy_report(
        self,
        company_slug: str,
        company_input: CompanyInput,
        research: ResearchOutput,
        synthesis: SynthesisOutput,
    ) -> str:
        """
        Generate final AI strategy report.

        Comprehensive document combining all deliverables into
        a professional consulting report.

        Args:
            company_slug: URL-safe company name.
            company_input: Original company input.
            research: Research output.
            synthesis: Synthesis output.

        Returns:
            Path to generated DOCX file.
        """
        doc = Document()

        # Set up custom styles
        self._setup_styles(doc)

        # Title page
        self._add_title_page(
            doc,
            company_input.name,
            "AI Strategy & Implementation Report",
            datetime.now().strftime("%B %Y"),
            logo_path=company_input.logo_path,
        )

        # Table of contents placeholder
        self._add_table_of_contents(doc)

        # Executive Summary
        self._add_section(doc, "Executive Summary", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "08_executive_summary", "Overview", 2)

        # Where You Stand Today
        self._add_section(doc, "Where You Stand Today", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "01_tools_audit", "Your Tool Stack & AI Opportunities", 2)

        # Where You're Losing Money
        self._add_section(doc, "Where You're Losing Money", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "02_daily_pain_points", "Time-Wasters & Bottlenecks", 2)

        # Your AI Readiness
        self._add_section(doc, "Your AI Readiness", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "05_readiness_assessment", "Readiness Assessment", 2)

        # What To Do First
        self._add_section(doc, "What To Do First", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "03_action_plan", "Priority Actions", 2)

        # Your Week-by-Week Plan
        self._add_section(doc, "Your Week-by-Week Plan", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "04_simple_roadmap", "Implementation Timeline", 2)

        # What It Costs & What You Save
        self._add_section(doc, "What It Costs & What You Save", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "06_roi_snapshot", "ROI Analysis", 2)

        # Putting It All Together
        self._add_section(doc, "Putting It All Together", 1)
        self._add_subsection_from_deliverable(doc, synthesis, "07_closing", "Summary", 2)

        # Save document
        output_path = self._get_output_path(company_slug, "final_strategy_report.docx")
        doc.save(output_path)

        return str(output_path)

    def _get_output_path(self, company_slug: str, filename: str) -> Path:
        """Get output path for a document."""
        output_dir = self.output_dir / company_slug / "documents"
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir / filename

    def _setup_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        # Get existing styles
        styles = doc.styles

        # Customize Heading 1
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = 'Arial'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(31, 78, 121)
            h1.paragraph_format.space_before = Pt(24)
            h1.paragraph_format.space_after = Pt(8)

        # Customize Heading 2
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = 'Arial'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(68, 114, 196)
            h2.paragraph_format.space_before = Pt(18)
            h2.paragraph_format.space_after = Pt(10)

        # Customize Heading 3
        if 'Heading 3' in styles:
            h3 = styles['Heading 3']
            h3.font.name = 'Arial'
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.font.color.rgb = RGBColor(68, 114, 196)
            h3.paragraph_format.space_before = Pt(12)
            h3.paragraph_format.space_after = Pt(8)

        # Customize Normal
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)
            normal.font.color.rgb = RGBColor(64, 64, 64)
            normal.paragraph_format.space_before = Pt(8)
            normal.paragraph_format.space_after = Pt(6)

        # Customize List Bullet
        if 'List Bullet' in styles:
            lb = styles['List Bullet']
            lb.paragraph_format.space_before = Pt(12)
            lb.paragraph_format.space_after = Pt(10)

        # Customize List Number
        if 'List Number' in styles:
            ln = styles['List Number']
            ln.paragraph_format.space_before = Pt(12)
            ln.paragraph_format.space_after = Pt(10)

        # Set up TOC 1 style with right-aligned tab stop and dot leader
        # 9360 twips = 6.5 inches (standard letter page text width)
        try:
            toc1 = styles['TOC 1']
        except KeyError:
            toc1 = styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
        pPr = toc1.element.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:leader'), 'none')
        tab_el.set(qn('w:pos'), '8500')
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

    def _add_title_page(
        self,
        doc: Document,
        company_name: str,
        title: str,
        subtitle: str,
        logo_path: Optional[str] = None,
    ) -> None:
        """Add title page to document."""
        # Add spacing at top (less if logo is present)
        spacing_count = 4 if (logo_path and Path(logo_path).exists()) else 8
        for _ in range(spacing_count):
            doc.add_paragraph()

        # Add company logo if provided
        if logo_path and Path(logo_path).exists():
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(str(logo_path), width=Inches(2.0))
            doc.add_paragraph()  # spacing after logo

        # Company name
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company_para.add_run(company_name)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)

        doc.add_paragraph()

        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(64, 64, 64)

        doc.add_paragraph()

        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Auditor branding footer — keep logo + text together on one page
        if _AUDITOR_LOGO.exists():
            # Push footer toward bottom of page (smaller spacing to prevent overflow)
            for _ in range(3):
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.line_spacing = Pt(14)

            # Auditor logo
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.keep_with_next = True
            run = logo_para.add_run()
            run.add_picture(str(_AUDITOR_LOGO), width=Inches(1.0))

            # "Prepared by" line
            prep_para = doc.add_paragraph()
            prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prep_para.paragraph_format.keep_with_next = True
            prep_para.paragraph_format.space_before = Pt(4)
            prep_para.paragraph_format.space_after = Pt(0)
            run = prep_para.add_run("Prepared by Early AIdopters")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(128, 128, 128)

            # Website line
            web_para = doc.add_paragraph()
            web_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            web_para.paragraph_format.space_before = Pt(0)
            web_para.paragraph_format.space_after = Pt(0)
            run = web_para.add_run("https://www.skool.com/earlyaidopters/about")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Add page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents using Word field codes (auto-generates on open)."""
        doc.add_heading("Table of Contents", level=1)

        paragraph = doc.add_paragraph()

        # Begin field
        run_begin = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        fld_begin.set(qn('w:dirty'), 'true')
        run_begin._r.append(fld_begin)

        # Field instruction: TOC levels 1-3, hyperlinked
        run_instr = paragraph.add_run()
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = ' TOC \\o "1-1" \\h \\z \\u '
        run_instr._r.append(instr)

        # Separate
        run_sep = paragraph.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fld_sep)

        # End field
        run_end = paragraph.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fld_end)

        doc.add_page_break()

    def _add_section(self, doc: Document, title: str, level: int) -> None:
        """Add a section heading."""
        doc.add_heading(title, level=level)

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """Add a paragraph with optional formatting."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

    def _add_bullet_list(self, doc: Document, items: List[str]) -> None:
        """Add a bullet list."""
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    def _add_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]],
    ) -> None:
        """Add a table to the document."""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Bold + dark blue text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(31, 78, 121)
            # Light blue background (#D6E4F0) matching PDF
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D6E4F0')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(value)


    def _add_subsection_from_deliverable(
        self,
        doc: Document,
        synthesis: SynthesisOutput,
        deliverable_id: str,
        title: str,
        level: int,
    ) -> None:
        """Add subsection from a deliverable's content."""
        if deliverable_id not in synthesis.deliverables:
            doc.add_paragraph("[Content not available]")
            return

        content = synthesis.deliverables[deliverable_id].content
        if not content:
            doc.add_paragraph("[Content not available]")
            return

        # Strip a leading H1 — the generator already added the section heading,
        # and some deliverables start with `# Title` while others don't.
        content = re.sub(r'^\s*#[^\n]*\n?', '', content, count=1)

        # Parse markdown content and convert to Word
        self._convert_markdown_to_docx(doc, content)

    def _convert_markdown_to_docx(self, doc: Document, markdown_content: str) -> None:
        """Convert markdown content to Word format."""
        lines = markdown_content.split('\n')
        in_code_block = False
        in_table = False
        in_yaml_front_matter = False
        table_lines = []
        last_para = None  # Track last paragraph for spacing adjustment before tables
        prev_was_numbered = False

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Handle YAML front matter (at the very start)
            if i == 0 and stripped == '---':
                in_yaml_front_matter = True
                i += 1
                continue

            if in_yaml_front_matter:
                if stripped == '---':
                    in_yaml_front_matter = False
                i += 1
                continue

            # Skip horizontal rules (---, ***, ___) - these cause the dashed lines issue
            if re.match(r'^[-*_]{3,}\s*$', stripped):
                i += 1
                continue

            # Handle code blocks
            if stripped.startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue

            if in_code_block:
                # Add as monospace (but limit line length)
                para = doc.add_paragraph()
                code_line = line[:200] if len(line) > 200 else line  # Truncate very long lines
                run = para.add_run(code_line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                i += 1
                continue

            # Handle tables
            if stripped.startswith('|') and stripped.endswith('|'):
                # Skip lines that are too long (corrupted tables)
                if len(stripped) > 1000:
                    # If we were in a table, end it and skip this corrupted line
                    if in_table and table_lines:
                        self._process_markdown_table(doc, table_lines)
                        table_lines = []
                    in_table = False
                    i += 1
                    continue
                in_table = True
                # Adjust gap on preceding paragraph before table
                if last_para is not None:
                    last_para.paragraph_format.space_after = Pt(12)
                table_lines.append(stripped)
                i += 1
                continue
            elif in_table:
                # End of table, process it
                if table_lines:
                    self._process_markdown_table(doc, table_lines)
                    table_lines = []
                in_table = False
                # Don't increment i - process this line normally

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Handle headings (but skip top-level as we add our own sections)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                # Remove any remaining markdown formatting from heading
                text = self._clean_markdown_text(text)
                # Map markdown H1 -> Word Heading 2 (H1 reserved for doc sections)
                word_level = 2 if level == 1 else min(level, 4)
                last_para = doc.add_heading(text, level=word_level)
                prev_was_numbered = False
                i += 1
                continue

            # Handle bullet points
            bullet_match = re.match(r'^[-*+]\s+(.+)$', stripped)
            if bullet_match:
                text = self._clean_markdown_text(bullet_match.group(1))
                if prev_was_numbered:
                    # Sub-item under numbered list — render as indented bullet
                    para = self._add_formatted_paragraph(doc, text, style='List Bullet')
                    para.paragraph_format.left_indent = Inches(0.75)
                else:
                    self._add_formatted_paragraph(doc, text, style='List Bullet')
                i += 1
                continue

            # Handle numbered lists
            numbered_match = re.match(r'^\d+\.\s+(.+)$', stripped)
            if numbered_match:
                if not prev_was_numbered:
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(6)
                text = self._clean_markdown_text(numbered_match.group(1))
                self._add_formatted_paragraph(doc, text, style='List Number')
                prev_was_numbered = True
                i += 1
                continue

            # Regular paragraph - add with inline formatting
            if stripped:
                if prev_was_numbered:
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(6)
                prev_was_numbered = False
                # Skip excessively long lines (corrupted data)
                if len(stripped) > 2000:
                    i += 1
                    continue
                # Skip lines that are mostly dashes (table separators that weren't caught)
                if stripped.count('-') > 50 and stripped.count('-') / len(stripped) > 0.5:
                    i += 1
                    continue
                last_para = self._add_formatted_paragraph(doc, stripped)
                # Bold-only lines (e.g. **Key Strengths:**) act as pseudo-headings — add space above
                if re.match(r'^\*\*.+\*\*:?\s*$', stripped):
                    last_para.paragraph_format.space_before = Pt(12)

            i += 1

        # Process any remaining table
        if table_lines:
            self._process_markdown_table(doc, table_lines)

    def _clean_markdown_text(self, text: str) -> str:
        """Remove markdown formatting markers from text."""
        # Remove markdown links but keep the text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Don't remove bold/italic markers here - we'll handle them in _add_formatted_paragraph
        return text.strip()

    def _add_formatted_paragraph(self, doc: Document, text: str, style: Optional[str] = None):
        """Add a paragraph with inline formatting (bold, italic, code). Returns the paragraph."""
        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        # Parse inline formatting
        # Pattern to match **bold**, *italic*, and `code`
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Regular text
                para.add_run(part)

        return para

    def _process_markdown_table(self, doc: Document, table_lines: List[str]) -> None:
        """Process markdown table lines into a Word table."""
        if len(table_lines) < 3:
            return

        # Check if table is corrupted (any line is excessively long - indicates garbage data)
        for line in table_lines:
            if len(line) > 1000:  # Normal tables shouldn't have lines this long
                # Skip this entire corrupted table
                return

        # Parse header
        raw_headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
        # Clean headers and truncate if too long
        headers = [self._clean_table_cell(h, max_length=100) for h in raw_headers]

        # Skip if headers look corrupted (all dashes or very long)
        if all(h.startswith('-') or len(h) > 200 for h in raw_headers):
            return

        # Check if any header cell is excessively long (indicates corruption)
        if any(len(h) > 500 for h in raw_headers):
            return

        # Parse rows (skip separator line at index 1)
        rows = []
        for line in table_lines[2:]:
            # Skip separator lines (contain only dashes and colons)
            if re.match(r'^[\|\s:\-]+$', line):
                continue

            # Skip lines that are mostly dashes (corrupted data)
            dash_count = line.count('-')
            if dash_count > 100:  # Too many dashes indicates garbage
                continue

            cells = [c.strip() for c in line.strip('|').split('|')]
            if len(cells) == len(headers):
                # Check if any cell is excessively long (indicates corruption)
                if any(len(c) > 2000 for c in cells):
                    continue

                # Clean and truncate each cell
                clean_cells = [self._clean_table_cell(c, max_length=500) for c in cells]
                # Skip rows that look corrupted (mostly dashes)
                if not all(c.startswith('-') for c in cells if c):
                    # Also skip if cleaned cells are mostly empty or dashes
                    if any(len(c) > 3 and not c.replace('-', '').replace('.', '').strip() == '' for c in clean_cells):
                        rows.append(clean_cells)

        if headers and rows:
            if self._is_action_table(raw_headers):
                self._add_action_cards(doc, headers, rows)
            elif self._is_roadmap_table(raw_headers):
                self._add_roadmap_cards(doc, headers, rows)
            else:
                self._add_table(doc, headers, rows)

    @staticmethod
    def _is_action_table(raw_headers: List[str]) -> bool:
        """Return True if this is the 7-column action plan table."""
        return len(raw_headers) == 7 and raw_headers[0].strip() == "#"

    @staticmethod
    def _is_roadmap_table(raw_headers: List[str]) -> bool:
        """Return True if this is the 5-column roadmap table."""
        return len(raw_headers) == 5 and raw_headers[0].strip() == "#"

    def _add_action_cards(self, doc: Document, headers: List[str], rows: List[List[str]]) -> None:
        """Render 7-column action table rows as bordered cards."""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)
        for row_data in rows:
            # row_data indices: 0=#  1=Action  2=Tool(Price)  3=Time to Set Up
            #                   4=Weekly Time Saved  5=What You Get  6=First Step Today

            card = doc.add_table(rows=1, cols=1)
            card.style = 'Table Grid'
            card.alignment = WD_TABLE_ALIGNMENT.CENTER

            cell = card.rows[0].cells[0]

            # Background shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F9FC')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

            # Clear default paragraph
            cell.paragraphs[0].clear()

            # Title: "1. Action text"
            title_para = cell.paragraphs[0]
            title_run = title_para.add_run(f"{row_data[0]}. {row_data[1]}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(31, 78, 121)
            title_para.paragraph_format.space_after = Pt(6)

            # Tool line
            tool_para = cell.add_paragraph()
            label_run = tool_para.add_run("Tool: ")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(64, 64, 64)
            value_run = tool_para.add_run(row_data[2])
            value_run.font.size = Pt(10)
            value_run.font.color.rgb = RGBColor(64, 64, 64)
            tool_para.paragraph_format.space_before = Pt(4)
            tool_para.paragraph_format.space_after = Pt(2)

            # Setup + Saves on one line
            detail_para = cell.add_paragraph()
            setup_label = detail_para.add_run("Setup: ")
            setup_label.bold = True
            setup_label.font.size = Pt(10)
            setup_label.font.color.rgb = RGBColor(64, 64, 64)
            setup_val = detail_para.add_run(row_data[3])
            setup_val.font.size = Pt(10)
            setup_val.font.color.rgb = RGBColor(64, 64, 64)
            detail_para.add_run("     ")
            saves_label = detail_para.add_run("Saves: ")
            saves_label.bold = True
            saves_label.font.size = Pt(10)
            saves_label.font.color.rgb = RGBColor(64, 64, 64)
            saves_val = detail_para.add_run(row_data[4])
            saves_val.font.size = Pt(10)
            saves_val.font.color.rgb = RGBColor(64, 64, 64)
            detail_para.paragraph_format.space_before = Pt(2)
            detail_para.paragraph_format.space_after = Pt(6)

            # What you get
            get_para = cell.add_paragraph()
            get_label = get_para.add_run("What you get: ")
            get_label.bold = True
            get_label.font.size = Pt(10)
            get_label.font.color.rgb = RGBColor(64, 64, 64)
            get_val = get_para.add_run(row_data[5])
            get_val.font.size = Pt(10)
            get_val.font.color.rgb = RGBColor(64, 64, 64)
            get_para.paragraph_format.space_before = Pt(2)
            get_para.paragraph_format.space_after = Pt(4)

            # First step today
            step_para = cell.add_paragraph()
            step_label = step_para.add_run("First step today: ")
            step_label.bold = True
            step_label.font.size = Pt(10)
            step_label.font.color.rgb = RGBColor(64, 64, 64)
            step_val = step_para.add_run(row_data[6])
            step_val.font.size = Pt(10)
            step_val.font.color.rgb = RGBColor(64, 64, 64)
            step_para.paragraph_format.space_before = Pt(2)
            step_para.paragraph_format.space_after = Pt(12)

    def _add_roadmap_cards(self, doc: Document, headers: List[str], rows: List[List[str]]) -> None:
        """Render 5-column roadmap table rows as simple bordered cards."""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)
        for row_data in rows:
            # row_data indices: 0=#  1=Action  2=Tool(Price)  3=Time  4=What To Do

            card = doc.add_table(rows=1, cols=1)
            card.style = 'Table Grid'
            card.alignment = WD_TABLE_ALIGNMENT.CENTER

            cell = card.rows[0].cells[0]

            # Background shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F9FC')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

            # Clear default paragraph
            cell.paragraphs[0].clear()

            # Title: "1. Action text"
            title_para = cell.paragraphs[0]
            title_run = title_para.add_run(f"{row_data[0]}. {row_data[1]}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(31, 78, 121)
            title_para.paragraph_format.space_after = Pt(4)

            # Tool + Time on one line
            detail_para = cell.add_paragraph()
            tool_label = detail_para.add_run("Tool: ")
            tool_label.bold = True
            tool_label.font.size = Pt(10)
            tool_label.font.color.rgb = RGBColor(64, 64, 64)
            tool_val = detail_para.add_run(row_data[2])
            tool_val.font.size = Pt(10)
            tool_val.font.color.rgb = RGBColor(64, 64, 64)
            detail_para.add_run("     ")
            time_label = detail_para.add_run("Time: ")
            time_label.bold = True
            time_label.font.size = Pt(10)
            time_label.font.color.rgb = RGBColor(64, 64, 64)
            time_val = detail_para.add_run(row_data[3])
            time_val.font.size = Pt(10)
            time_val.font.color.rgb = RGBColor(64, 64, 64)
            detail_para.paragraph_format.space_before = Pt(2)
            detail_para.paragraph_format.space_after = Pt(4)

            # What To Do
            do_para = cell.add_paragraph()
            do_label = do_para.add_run("What to do: ")
            do_label.bold = True
            do_label.font.size = Pt(10)
            do_label.font.color.rgb = RGBColor(64, 64, 64)
            do_val = do_para.add_run(row_data[4])
            do_val.font.size = Pt(10)
            do_val.font.color.rgb = RGBColor(64, 64, 64)
            do_para.paragraph_format.space_before = Pt(2)
            do_para.paragraph_format.space_after = Pt(12)

    def _clean_table_cell(self, text: str, max_length: int = 500) -> str:
        """Clean and truncate table cell content."""
        # Remove markdown formatting
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
        text = re.sub(r'\*([^*]+)\*', r'\1', text)  # *italic*
        text = re.sub(r'`([^`]+)`', r'\1', text)  # `code`
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # [link](url)

        # Truncate if too long
        if len(text) > max_length:
            text = text[:max_length-3] + '...'

        return text.strip()




# Convenience function

def generate_strategy_report(
    company_slug: str,
    company_input: CompanyInput,
    research: ResearchOutput,
    synthesis: SynthesisOutput,
    output_dir: Optional[Path] = None,
) -> str:
    """Generate final AI strategy report."""
    generator = DocxGenerator(output_dir=output_dir)
    return generator.generate_strategy_report(
        company_slug, company_input, research, synthesis
    )
